"""
Document processing utilities for Betty AI Assistant.

This module provides reusable functions for extracting text from various
document formats and processing them for use in the RAG system.
"""

import io
import re
from typing import List, Optional
import PyPDF2
import docx
import streamlit as st
import tiktoken
from config.settings import AppConfig
try:
    import nltk
    from nltk.tokenize import sent_tokenize
    NLTK_AVAILABLE = True
except ImportError:
    NLTK_AVAILABLE = False


class DocumentProcessor:
    """Document processing utilities with improved error handling."""
    
    def __init__(self, tokenizer_model: str = None):
        """Initialize the document processor.
        
        Args:
            tokenizer_model: The tokenizer model to use for chunking.
        """
        self.tokenizer = tiktoken.get_encoding(
            tokenizer_model or AppConfig.TOKENIZER_MODEL
        )
        self._ensure_nltk_data()
    
    def _ensure_nltk_data(self):
        """Ensure required NLTK data is available."""
        if NLTK_AVAILABLE:
            try:
                nltk.data.find('tokenizers/punkt')
            except LookupError:
                try:
                    nltk.download('punkt', quiet=True)
                except Exception:
                    pass  # Fail silently if download fails
    
    def extract_text_from_pdf(self, file: io.BytesIO) -> str:
        """Extract text from an in-memory PDF file.
        
        Args:
            file: BytesIO object containing PDF data.
            
        Returns:
            Extracted text as string, empty string if extraction fails.
        """
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            pages_text = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        pages_text.append(text)
                except Exception as e:
                    st.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
                    
            return "\n".join(pages_text)
            
        except Exception as e:
            st.error(f"Error reading PDF file: {e}")
            return ""
    
    def extract_text_from_docx(self, file: io.BytesIO) -> str:
        """Extract text from an in-memory DOCX file with structure preservation.
        
        Args:
            file: BytesIO object containing DOCX data.
            
        Returns:
            Extracted text with preserved structure, empty string if extraction fails.
        """
        try:
            doc = docx.Document(file)
            content_parts = []
            
            # Process paragraphs with style information
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Preserve heading structure
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    if level.isdigit():
                        heading_level = int(level)
                        prefix = "#" * min(heading_level, 6)
                        content_parts.append(f"\n{prefix} {text}\n")
                    else:
                        content_parts.append(f"\n## {text}\n")
                # Preserve list structure
                elif para.style.name.startswith('List'):
                    content_parts.append(f"• {text}")
                # Regular paragraphs
                else:
                    content_parts.append(text)
            
            # Process tables
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_cells.append(cell_text)
                    if row_cells:
                        table_content.append(" | ".join(row_cells))
                
                if table_content:
                    content_parts.append("\n--- Table ---")
                    content_parts.extend(table_content)
                    content_parts.append("--- End Table ---\n")
            
            return "\n".join(content_parts)
            
        except Exception as e:
            st.error(f"Error reading DOCX file: {e}")
            return ""
    
    def extract_text_from_txt(self, file: io.BytesIO) -> str:
        """Extract text from a plain text file.
        
        Args:
            file: BytesIO object containing text data.
            
        Returns:
            Extracted text as string, empty string if extraction fails.
        """
        try:
            return file.read().decode('utf-8')
        except UnicodeDecodeError:
            try:
                file.seek(0)
                return file.read().decode('latin-1')
            except Exception as e:
                st.error(f"Error reading text file: {e}")
                return ""
        except Exception as e:
            st.error(f"Error reading text file: {e}")
            return ""
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text.
            
        Returns:
            Cleaned text with normalized spacing and formatting.
        """
        if not text:
            return ""
        
        # Fix common formatting issues
        text = re.sub(r'([.,])([a-zA-Z])', r'\1 \2', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Clean up line breaks and spacing
        lines = (line.strip() for line in text.splitlines())
        return "\n".join(line for line in lines if line)
    
    def chunk_text(
        self, 
        text: str, 
        chunk_size: int = None, 
        overlap: int = None
    ) -> List[str]:
        """Split text into overlapping chunks based on token count.
        
        Args:
            text: Text to chunk.
            chunk_size: Size of each chunk in tokens.
            overlap: Number of overlapping tokens between chunks.
            
        Returns:
            List of text chunks.
        """
        chunk_size = chunk_size or AppConfig.CHUNK_SIZE
        overlap = overlap or AppConfig.CHUNK_OVERLAP
        
        # Validate parameters
        if overlap >= chunk_size:
            st.warning(f"Overlap ({overlap}) must be less than chunk size ({chunk_size})")
            overlap = chunk_size // 4
        
        try:
            tokens = self.tokenizer.encode(text)
            chunks = []
            
            for i in range(0, len(tokens), chunk_size - overlap):
                chunk_tokens = tokens[i:i + chunk_size]
                chunk_text = self.tokenizer.decode(chunk_tokens)
                
                if chunk_text.strip():
                    chunks.append(chunk_text)
            
            return chunks
            
        except Exception as e:
            st.error(f"Error chunking text: {e}")
            return [text]  # Return original text as fallback
    
    def semantic_chunk_text(
        self, 
        text: str, 
        chunk_size: int = None, 
        overlap: int = None
    ) -> List[str]:
        """Split text into semantic chunks using sentence boundaries.
        
        Args:
            text: Text to chunk.
            chunk_size: Size of each chunk in tokens.
            overlap: Number of overlapping tokens between chunks.
            
        Returns:
            List of semantic text chunks.
        """
        if not AppConfig.USE_SEMANTIC_CHUNKING or not NLTK_AVAILABLE:
            return self.chunk_text(text, chunk_size, overlap)
        
        chunk_size = chunk_size or AppConfig.CHUNK_SIZE
        overlap = overlap or AppConfig.CHUNK_OVERLAP
        
        try:
            # Split into sentences using NLTK
            sentences = sent_tokenize(text)
            chunks = []
            current_chunk = ""
            current_tokens = 0
            
            for sentence in sentences:
                sentence_tokens = len(self.tokenizer.encode(sentence))
                
                # If adding this sentence would exceed chunk size and we have content
                if current_tokens + sentence_tokens > chunk_size and current_chunk:
                    chunks.append(current_chunk.strip())
                    
                    # Create overlap by keeping last few sentences
                    if overlap > 0:
                        overlap_sentences = current_chunk.split('. ')
                        overlap_text = ""
                        overlap_tokens = 0
                        
                        # Add sentences from the end until we reach overlap limit
                        for i in range(len(overlap_sentences) - 1, -1, -1):
                            sentence_text = overlap_sentences[i] + '. '
                            sentence_token_count = len(self.tokenizer.encode(sentence_text))
                            
                            if overlap_tokens + sentence_token_count <= overlap:
                                overlap_text = sentence_text + overlap_text
                                overlap_tokens += sentence_token_count
                            else:
                                break
                        
                        current_chunk = overlap_text + sentence + '. '
                        current_tokens = len(self.tokenizer.encode(current_chunk))
                    else:
                        current_chunk = sentence + '. '
                        current_tokens = sentence_tokens
                else:
                    current_chunk += sentence + '. '
                    current_tokens += sentence_tokens
            
            # Add final chunk if there's content
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            
            return chunks if chunks else [text]
            
        except Exception as e:
            st.error(f"Error in semantic chunking: {e}")
            # Fallback to regular chunking
            return self.chunk_text(text, chunk_size, overlap)
    
    def get_file_type(self, filename: str) -> Optional[str]:
        """Determine file type from filename.
        
        Args:
            filename: Name of the file.
            
        Returns:
            File type string or None if unsupported.
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return 'pdf'
        elif filename_lower.endswith('.docx'):
            return 'docx'
        elif filename_lower.endswith('.txt'):
            return 'txt'
        else:
            return None
    
    def process_uploaded_file(self, uploaded_file) -> str:
        """Process an uploaded file and extract text.
        
        Args:
            uploaded_file: Streamlit uploaded file object.
            
        Returns:
            Extracted and cleaned text.
        """
        if not uploaded_file:
            return ""
        
        # Check file size
        if uploaded_file.size > AppConfig.MAX_FILE_SIZE_MB * 1024 * 1024:
            st.error(f"File {uploaded_file.name} is too large "
                    f"(max {AppConfig.MAX_FILE_SIZE_MB}MB)")
            return ""
        
        file_type = self.get_file_type(uploaded_file.name)
        if not file_type:
            st.warning(f"Unsupported file type: {uploaded_file.name}")
            return ""
        
        try:
            file_bytes = io.BytesIO(uploaded_file.getvalue())
            
            if file_type == 'pdf':
                text = self.extract_text_from_pdf(file_bytes)
            elif file_type == 'docx':
                text = self.extract_text_from_docx(file_bytes)
            elif file_type == 'txt':
                text = self.extract_text_from_txt(file_bytes)
            else:
                return ""
            
            return self.clean_text(text)
            
        except Exception as e:
            st.error(f"Error processing file {uploaded_file.name}: {e}")
            return ""


# Create a global instance for easy importing
document_processor = DocumentProcessor()